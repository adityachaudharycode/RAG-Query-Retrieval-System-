"""
Document processing service for extracting text from PDFs and DOCX files
"""
import os
import re
import tempfile
from typing import List, Dict, Any
import requests
import fitz  # PyMuPDF
from docx import Document
from loguru import logger
from urllib.parse import urlparse

from config import settings
from models.schemas import DocumentChunk


class DocumentProcessor:
    """Service for processing various document formats"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
        # Optimize chunk size for faster processing (larger chunks = fewer embeddings)
        self.chunk_size = settings.chunk_size * 2  # Double the chunk size
        self.chunk_overlap = max(50, settings.chunk_overlap // 2)  # Reduce overlap
    
    async def process_document(self, document_url: str) -> str:
        """
        Download and process a document from URL (supports Google Drive, direct URLs)

        Args:
            document_url: URL to the document

        Returns:
            Extracted text content
        """
        try:
            logger.info(f"Processing document from URL: {document_url}")

            # Handle Google Drive URLs
            processed_url = self._process_google_drive_url(document_url)

            # Download the document with proper headers
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }

            response = requests.get(processed_url, headers=headers, timeout=30, allow_redirects=True)
            response.raise_for_status()

            # Determine file type from URL, content-type, or file signature
            file_extension = self._get_file_extension(document_url, response.headers, response.content)

            # Save to temporary file
            with tempfile.NamedTemporaryFile(suffix=file_extension, delete=False) as temp_file:
                temp_file.write(response.content)
                temp_file_path = temp_file.name

            try:
                # Extract text based on file type
                if file_extension.lower() == '.pdf':
                    text = self._extract_text_from_pdf(temp_file_path)
                elif file_extension.lower() in ['.docx', '.doc']:
                    text = self._extract_text_from_docx(temp_file_path)
                else:
                    raise ValueError(f"Unsupported file format: {file_extension}")

                logger.info(f"Successfully extracted {len(text)} characters from document")
                return text

            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)

        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise
    
    def _process_google_drive_url(self, url: str) -> str:
        """Convert Google Drive sharing URL to direct download URL"""
        if 'drive.google.com' in url:
            # Handle different Google Drive URL formats
            if '/file/d/' in url:
                # Extract file ID from sharing URL
                file_id = url.split('/file/d/')[1].split('/')[0]
                return f"https://drive.google.com/uc?export=download&id={file_id}"
            elif 'id=' in url:
                # Already in download format
                return url
        return url

    def _get_file_extension(self, url: str, headers: Dict[str, str], content: bytes = None) -> str:
        """Determine file extension from URL, content-type header, or file signature"""

        # Try to get extension from URL
        parsed_url = urlparse(url)
        path = parsed_url.path
        if path:
            _, ext = os.path.splitext(path)
            if ext.lower() in self.supported_formats:
                return ext

        # Try to get from content-type header
        content_type = headers.get('content-type', '').lower()
        if 'pdf' in content_type:
            return '.pdf'
        elif 'word' in content_type or 'officedocument' in content_type or 'msword' in content_type:
            return '.docx'

        # Try to detect from file signature (magic bytes) if content is available
        if content:
            file_ext = self._detect_file_type_from_content(content)
            if file_ext:
                return file_ext

        # For Google Drive URLs, try to detect from content disposition
        content_disposition = headers.get('content-disposition', '')
        if content_disposition:
            if '.pdf' in content_disposition.lower():
                return '.pdf'
            elif '.docx' in content_disposition.lower() or '.doc' in content_disposition.lower():
                return '.docx'

        # Default to PDF if uncertain
        return '.pdf'

    def _detect_file_type_from_content(self, content: bytes) -> str:
        """Detect file type from file signature (magic bytes)"""
        if not content or len(content) < 8:
            return None

        # PDF signature
        if content.startswith(b'%PDF'):
            return '.pdf'

        # DOCX signature (ZIP-based format)
        if content.startswith(b'PK\x03\x04') or content.startswith(b'PK\x05\x06') or content.startswith(b'PK\x07\x08'):
            # Check if it's a DOCX by looking for specific content
            if b'word/' in content[:1024] or b'document.xml' in content[:2048]:
                return '.docx'

        # DOC signature
        if content.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
            return '.doc'

        return None
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using PyMuPDF"""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                text += f"\n--- Page {page_num + 1} ---\n{page_text}"
            
            doc.close()
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\"\'\/\%\$\@\#]', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()
    
    async def chunk_text(self, text: str) -> List[DocumentChunk]:
        """
        Split text into chunks for vector embedding
        
        Args:
            text: Input text to chunk
            
        Returns:
            List of DocumentChunk objects
        """
        try:
            chunks = []
            
            # Split text into sentences first
            sentences = self._split_into_sentences(text)
            
            current_chunk = ""
            chunk_id = 0
            
            for sentence in sentences:
                # Check if adding this sentence would exceed chunk size
                if len(current_chunk) + len(sentence) > self.chunk_size:
                    if current_chunk:
                        # Create chunk
                        chunk = DocumentChunk(
                            id=f"chunk_{chunk_id}",
                            content=current_chunk.strip(),
                            metadata={
                                "chunk_index": chunk_id,
                                "length": len(current_chunk),
                                "sentence_count": len(current_chunk.split('.'))
                            }
                        )
                        chunks.append(chunk)
                        chunk_id += 1
                        
                        # Start new chunk with overlap
                        if self.chunk_overlap > 0:
                            overlap_text = current_chunk[-self.chunk_overlap:]
                            current_chunk = overlap_text + " " + sentence
                        else:
                            current_chunk = sentence
                    else:
                        current_chunk = sentence
                else:
                    current_chunk += " " + sentence if current_chunk else sentence
            
            # Add the last chunk
            if current_chunk:
                chunk = DocumentChunk(
                    id=f"chunk_{chunk_id}",
                    content=current_chunk.strip(),
                    metadata={
                        "chunk_index": chunk_id,
                        "length": len(current_chunk),
                        "sentence_count": len(current_chunk.split('.'))
                    }
                )
                chunks.append(chunk)
            
            logger.info(f"Created {len(chunks)} chunks from text")
            return chunks
            
        except Exception as e:
            logger.error(f"Error chunking text: {str(e)}")
            raise
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences using simple regex"""
        # Simple sentence splitting - can be improved with NLTK
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
